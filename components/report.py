from docx import Document
from utils.weibull import weibull_parameters_summary
import pandas as pd
import streamlit as st

def generate_report(data, shape, scale):
    doc = Document()
    doc.add_heading('Weibull Distribution Analysis Report', 0)

    doc.add_heading('Weibull Parameters', level=1)
    summary = weibull_parameters_summary(shape, scale)
    for key, value in summary.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('Sample Data', level=1)
    doc.add_paragraph(str(data.head()))

    # Add more sections as needed

    # Save to a BytesIO object for download
    import io
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def download_report(report_path):
    with open(report_path, "rb") as f:
        st.download_button(
            label="Download Report",
            data=f,
            file_name="weibull_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )